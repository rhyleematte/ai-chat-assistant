import sys
from docx import Document

DOCX_PATH = r"c:\Users\Sofia\ai-client-assistant-1233\community_cafe_full_framework.docx"

def update_docx():
    try:
        doc = Document(DOCX_PATH)
        
        # Add a new section for Business Details
        doc.add_paragraph('---') # Separator
        p_title = doc.add_paragraph()
        p_title.add_run('Business Location & Operating Hours').bold = True
        p_title.add_run('\n')
        
        p = doc.add_paragraph()
        p.add_run('Location: ').bold = True
        p.add_run('Butuan City')
        
        p = doc.add_paragraph()
        p.add_run('Operating Hours: ').bold = True
        p.add_run('7:00 AM to 9:00 PM, Monday to Friday')
        
        doc.save(DOCX_PATH)
        print("Successfully updated community_cafe_full_framework.docx with location and hours.")
        
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    update_docx()
