# -*- coding: utf-8 -*-
import sys, io
if sys.stdout.encoding and sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')
"""
community_guidelines_ai.py
──────────────────────────
Strata Management Consultant — Community Guidelines AI
Reads community_cafe_full_framework.docx and uses it as the knowledge
base to answer community questions.

If a question is OUTSIDE the community guidelines, a seeder staff profile
(Community Liaison Officer) generates the suggested response and recommended
action instead.

Usage:
    python community_guidelines_ai.py
    python community_guidelines_ai.py --demo   # runs preset demo questions
"""

import os
import sys
import json
import textwrap
import argparse
from pathlib import Path
from typing import Optional

# ── Third-party ──────────────────────────────────────────────────────────────
try:
    from docx import Document
except ImportError:
    sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")

try:
    from google import genai
    from google.genai import types as genai_types
except ImportError:
    sys.exit("ERROR: google-genai not installed. Run: pip install google-genai")

try:
    from rich.console import Console
    from rich.panel import Panel
    from rich.table import Table
    from rich.markdown import Markdown
    from rich import box
    from rich.prompt import Prompt
    RICH = True
except ImportError:
    RICH = False

# ── Config ───────────────────────────────────────────────────────────────────
DOCX_PATH = Path(__file__).parent / "community_cafe_full_framework.docx"
ENV_PATH  = Path(__file__).parent / ".env"

# Models tried in order -- falls back if quota exceeded on a model
GEMINI_MODELS = [
    "models/gemini-2.0-flash",
    "models/gemini-2.0-flash-lite",
    "models/gemini-2.5-flash",
]

# ── Seeder profile ───────────────────────────────────────────────────────────
SEEDER_PROFILE = {
    "name":  "Maria Santos",
    "role":  "Community Liaison Officer",
    "email": "m.santos@strata-consultants.com.au",
    "bio":   (
        "Senior Community Liaison Officer with 8 years of experience in "
        "strata and body-corporate management. Specialist in resident "
        "relations, by-laws interpretation, and community engagement programs."
    ),
}

# ── Helpers ──────────────────────────────────────────────────────────────────

def load_env(path: Path) -> dict:
    """Parse a .env file into a dict."""
    env: dict = {}
    if not path.exists():
        return env
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, _, v = line.partition("=")
        env[k.strip()] = v.strip().strip('"').strip("'")
    return env


def get_gemini_key() -> str:
    env = load_env(ENV_PATH)
    key = (
        os.environ.get("GEMINI_API_KEY")
        or env.get("GEMINI_API_KEY")
        or env.get("VITE_GEMINI_API_KEY")
    )
    if not key:
        sys.exit(
            "ERROR: GEMINI_API_KEY not found. "
            "Set it in your .env file or environment."
        )
    return key


def extract_docx_text(path: Path) -> str:
    """Extract all readable text from the .docx file."""
    if not path.exists():
        sys.exit(f"ERROR: DOCX file not found: {path}")

    doc = Document(str(path))
    parts: list[str] = []

    # Paragraphs
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


# ── Core AI logic ─────────────────────────────────────────────────────────────

SYSTEM_PROMPT_TEMPLATE = """
You are the AI assistant for Strata Management Consultants, embedded in a
community management platform.

Your PRIMARY knowledge base is the following community framework document.
Read it carefully — it describes the community cafe guidelines, mission,
values, menu, pricing, and programs that operate within the strata property.

────────────────────────────────────────────────────
COMMUNITY FRAMEWORK DOCUMENT:
{guidelines}
────────────────────────────────────────────────────

ROLES:
  • When a question is ANSWERABLE using the community guidelines above,
    respond as the AI assistant and set "within_guidelines": true.
  • When a question is OUTSIDE or BEYOND the community guidelines
    (e.g., legal matters, maintenance, by-laws, levy disputes, governance,
    safety, incidents, complaints about residents, external services), set
    "within_guidelines": false and hand off to the seeder staff profile.

SEEDER STAFF PROFILE (used when outside guidelines):
  Name:  {staff_name}
  Role:  {staff_role}
  Email: {staff_email}
  Bio:   {staff_bio}

RESPONSE FORMAT — return ONLY valid JSON, no markdown fences, no extra text:
{{
  "within_guidelines": true | false,
  "category": "one of: menu | pricing | hours | mission | values | events | programs | complaint | maintenance | legal | billing | governance | other",
  "confidence": 0.0 to 1.0,
  "suggested_response": "Professional reply to the resident/community member (2–4 short paragraphs, warm tone, address by name if known)",
  "recommended_action": "Short internal next-step for staff",
  "handled_by": "AI Assistant" | "{staff_name} — {staff_role}",
  "escalate_to_email": null | "{staff_email}",
  "reasoning": "One sentence explaining why this was or wasn't within guidelines"
}}
""".strip()


def build_system_prompt(guidelines_text: str) -> str:
    return SYSTEM_PROMPT_TEMPLATE.format(
        guidelines=guidelines_text,
        staff_name=SEEDER_PROFILE["name"],
        staff_role=SEEDER_PROFILE["role"],
        staff_email=SEEDER_PROFILE["email"],
        staff_bio=SEEDER_PROFILE["bio"],
    )


def ask_ai(
    api_key: str,
    system_prompt: str,
    resident_name: str,
    question: str,
) -> dict:
    user_message = (
        f"Resident name: {resident_name}\n"
        f"Question: {question}"
    )
    full_prompt = system_prompt + "\n\n" + user_message

    last_err = None
    client = genai.Client(api_key=api_key)
    for model_name in GEMINI_MODELS:
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=full_prompt,
                config=genai_types.GenerateContentConfig(
                    response_mime_type="application/json",
                    temperature=0.3,
                    max_output_tokens=1024,
                ),
            )
            raw = (response.text or "{}").strip()
            # Strip markdown code fences if model wraps the response
            if raw.startswith("```"):
                raw = raw.split("\n", 1)[-1]
                raw = raw.rsplit("```", 1)[0].strip()
            result = json.loads(raw)
            result["_model_used"] = model_name
            return result
        except Exception as e:
            msg = str(e)
            last_err = msg
            if "429" in msg or "quota" in msg.lower() or "exhausted" in msg.lower() or "resource_exhausted" in msg.lower():
                continue   # try next model
            break          # non-quota error -- fail immediately

    # All models failed
    return {
        "within_guidelines": False,
        "category": "other",
        "confidence": 0.0,
        "suggested_response": (
            "We are temporarily unable to process your enquiry automatically. "
            "Please contact our team directly."
        ),
        "recommended_action": f"AI unavailable: {last_err[:200] if last_err else 'Unknown error'}",
        "handled_by": SEEDER_PROFILE["name"] + " -- " + SEEDER_PROFILE["role"],
        "escalate_to_email": SEEDER_PROFILE["email"],
        "reasoning": "Gemini quota exceeded on all available models.",
        "_model_used": "none",
    }


# ── Display ───────────────────────────────────────────────────────────────────

def print_result(result: dict, resident: str, question: str):
    if not RICH:
        # Plain fallback
        print("\n" + "=" * 60)
        print(f"QUESTION FROM: {resident}")
        print(f"QUESTION: {question}")
        print(f"WITHIN GUIDELINES: {result.get('within_guidelines')}")
        print(f"CATEGORY: {result.get('category')}")
        print(f"CONFIDENCE: {result.get('confidence', 0):.0%}")
        print(f"HANDLED BY: {result.get('handled_by')}")
        print(f"\nSUGGESTED RESPONSE:\n{result.get('suggested_response')}")
        print(f"\nRECOMMENDED ACTION: {result.get('recommended_action')}")
        if result.get("escalate_to_email"):
            print(f"ESCALATE TO: {result.get('escalate_to_email')}")
        print("=" * 60)
        return

    console = Console()

    within = result.get("within_guidelines", False)
    handled = result.get("handled_by", "Unknown")
    category = result.get("category", "other").upper()
    confidence = result.get("confidence", 0)
    conf_str = f"{confidence:.0%}"

    # Header panel
    status_color = "green" if within else "yellow"
    status_label = "✅ Within Guidelines" if within else "⚠️  Beyond Guidelines — Escalated"
    console.print(Panel(
        f"[bold]{status_label}[/bold]\n"
        f"Category: [cyan]{category}[/cyan]  |  "
        f"Confidence: [cyan]{conf_str}[/cyan]  |  "
        f"Handled by: [cyan]{handled}[/cyan]",
        title=f"[bold white]Enquiry from: {resident}[/bold white]",
        border_style=status_color,
        expand=False,
    ))

    # Question
    console.print(f"\n[bold]Question:[/bold] {question}\n")

    # Suggested response
    console.print(Panel(
        str(result.get("suggested_response") or "-"),
        title="[bold green]Suggested Response[/bold green]",
        border_style="green",
    ))

    # Recommended action
    console.print(Panel(
        str(result.get("recommended_action") or "-"),
        title="[bold blue]Recommended Action[/bold blue]",
        border_style="blue",
    ))

    # Escalation info (only when outside guidelines)
    if not within and result.get("escalate_to_email"):
        seeder = Table(box=box.ROUNDED, show_header=False, border_style="yellow")
        seeder.add_column("Field", style="bold yellow", width=18)
        seeder.add_column("Value")
        seeder.add_row("Staff Profile", SEEDER_PROFILE["name"])
        seeder.add_row("Role", SEEDER_PROFILE["role"])
        seeder.add_row("Email", SEEDER_PROFILE["email"])
        seeder.add_row("Escalate to", result.get("escalate_to_email", "-"))
        seeder.add_row("Bio", textwrap.fill(SEEDER_PROFILE["bio"], width=60))
        console.print(Panel(
            seeder,
            title="[bold yellow]Seeder Staff Profile -- Escalation[/bold yellow]",
            border_style="yellow",
        ))

    # Reasoning
    console.print(
        f"\n[dim]Reasoning: {result.get('reasoning', '')}[/dim]\n"
    )


# ── Demo questions ────────────────────────────────────────────────────────────

DEMO_QUESTIONS = [
    {
        "resident": "Ana Reyes",
        "question": "What's on the breakfast menu and how much does it cost?",
    },
    {
        "resident": "Mark Villanueva",
        "question": "Is there a student discount program at the cafe?",
    },
    {
        "resident": "Linda Cruz",
        "question": (
            "My neighbor is parking in my allocated space and blocking my car. "
            "What can I do about it?"
        ),
    },
    {
        "resident": "James Tan",
        "question": (
            "I received a levy notice that seems higher than what was agreed. "
            "Who handles billing disputes?"
        ),
    },
    {
        "resident": "Sofia Mendez",
        "question": "What drinks are available and what's the cheapest option?",
    },
    {
        "resident": "Roberto Lim",
        "question": (
            "There's a water leak coming from the unit above me. "
            "This is urgent — what should I do?"
        ),
    },
    {
        "resident": "Sofia",
        "question": "What are your operating hours and where are you located?",
    },
]


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Strata Management — Community Guidelines AI"
    )
    parser.add_argument(
        "--demo", action="store_true",
        help="Run preset demo questions instead of interactive mode"
    )
    args = parser.parse_args()

    if RICH:
        console = Console()
        console.print(Panel(
            "[bold white]Strata Management Consultants[/bold white]\n"
            "[cyan]Community Guidelines AI -- Powered by Gemini[/cyan]\n\n"
            f"Knowledge base: [yellow]{DOCX_PATH.name}[/yellow]\n"
            f"Models:         [yellow]{', '.join(GEMINI_MODELS)}[/yellow]\n"
            f"Seeder profile: [yellow]{SEEDER_PROFILE['name']} "
            f"({SEEDER_PROFILE['role']})[/yellow]",
            title="[bold green]AI System Ready[/bold green]",
            border_style="green",
        ))
    else:
        print("=== Strata Management — Community Guidelines AI ===")

    # Load docs
    if RICH:
        console.print("[dim]Loading community framework document...[/dim]")
    guidelines = extract_docx_text(DOCX_PATH)
    if RICH:
        console.print(
            f"[green]✓[/green] Loaded {len(guidelines):,} characters "
            f"from '{DOCX_PATH.name}'\n"
        )

    # Build Gemini client (key only -- client created per-call for model fallback)
    gemini_key = get_gemini_key()
    system_prompt = build_system_prompt(guidelines)

    if args.demo:
        # ── Demo mode ───────────────────────────────────────────────────────
        if RICH:
            console.print(
                f"[bold]Running {len(DEMO_QUESTIONS)} demo questions...[/bold]\n"
            )
        for i, q in enumerate(DEMO_QUESTIONS, 1):
            if RICH:
                console.rule(f"[bold]Demo Question {i} of {len(DEMO_QUESTIONS)}[/bold]")
            result = ask_ai(gemini_key, system_prompt, q["resident"], q["question"])
            print_result(result, q["resident"], q["question"])
    else:
        # ── Interactive mode ─────────────────────────────────────────────────
        if RICH:
            console.print(
                "[bold]Interactive mode[/bold] -- type [cyan]quit[/cyan] or "
                "[cyan]exit[/cyan] to stop, [cyan]demo[/cyan] to run demos.\n"
            )
        while True:
            try:
                if RICH:
                    resident = Prompt.ask("[bold cyan]Resident name[/bold cyan]")
                else:
                    resident = input("Resident name: ").strip()

                if resident.lower() in ("quit", "exit", "q"):
                    break
                if resident.lower() == "demo":
                    for i, q in enumerate(DEMO_QUESTIONS, 1):
                        if RICH:
                            console.rule(f"Demo {i} of {len(DEMO_QUESTIONS)}")
                        result = ask_ai(
                            gemini_key, system_prompt, q["resident"], q["question"]
                        )
                        print_result(result, q["resident"], q["question"])
                    continue

                if RICH:
                    question = Prompt.ask("[bold cyan]Question          [/bold cyan]")
                else:
                    question = input("Question: ").strip()

                if not question:
                    continue

                if RICH:
                    console.print("[dim]Analysing...[/dim]")
                result = ask_ai(gemini_key, system_prompt, resident, question)
                print_result(result, resident, question)

            except KeyboardInterrupt:
                if RICH:
                    console.print("\n[yellow]Exiting.[/yellow]")
                break


if __name__ == "__main__":
    main()
